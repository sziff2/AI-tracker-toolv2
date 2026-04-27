"""
Document Processing / Parsing Service (§7)

Responsibilities:
  - Classify document type via LLM
  - Extract text from PDFs / DOCX
  - Detect and extract tables
  - Split into sections
  - Persist parsed results as JSON and DocumentSection rows
"""

import json
import logging
import uuid
from pathlib import Path

import fitz  # pymupdf
import pdfplumber

from sqlalchemy.ext.asyncio import AsyncSession

from apps.api.models import Document, DocumentSection
from configs.settings import settings
from prompts import DOCUMENT_CLASSIFIER
from schemas import ClassifiedDocument
from services.llm_client import call_llm_json

logger = logging.getLogger(__name__)


# ─────────────────────────────────────────────────────────────────
# Text extraction
# ─────────────────────────────────────────────────────────────────

def extract_text_pymupdf(file_path: str) -> list[dict]:
    """Return a list of {page, text} dicts using PyMuPDF."""
    doc = fitz.open(file_path)
    pages = []
    for i, page in enumerate(doc):
        pages.append({"page": i + 1, "text": page.get_text("text")})
    doc.close()
    return pages


def extract_tables_pdfplumber(file_path: str) -> list[dict]:
    """Return a list of {page, tables} using pdfplumber."""
    results = []
    with pdfplumber.open(file_path) as pdf:
        for i, page in enumerate(pdf.pages):
            tables = page.extract_tables()
            if tables:
                results.append({"page": i + 1, "tables": tables})
    return results


def extract_text_docx(file_path: str) -> list[dict]:
    """Extract text from a .docx file."""
    from docx import Document as DocxDoc
    doc = DocxDoc(file_path)
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    return [{"page": 1, "text": "\n".join(paragraphs)}]


def extract_text_xlsx(file_path: str) -> list[dict]:
    """Extract text from a .xlsx workbook. One 'page' per sheet, with
    cells joined into tab-separated rows that preserve column alignment.
    Used for consensus packs that arrive as Bloomberg / VARA / Visible
    Alpha exports (Excel format).

    For .xls (legacy binary), openpyxl can't read it — caller gets
    empty text and the consensus extractor logs a soft failure rather
    than crashing. .xlsx covers ~95% of modern exports."""
    try:
        from openpyxl import load_workbook
    except ImportError:
        return [{"page": 1, "text": ""}]

    pages: list[dict] = []
    try:
        wb = load_workbook(file_path, read_only=True, data_only=True)
    except Exception as exc:
        # .xls (legacy), corrupt file, password-protected, etc.
        logger.warning("openpyxl failed on %s: %s", Path(file_path).name, str(exc)[:200])
        return [{"page": 1, "text": ""}]

    for i, sheet_name in enumerate(wb.sheetnames):
        sheet = wb[sheet_name]
        rows: list[str] = [f"### Sheet: {sheet_name}"]
        for row in sheet.iter_rows(values_only=True):
            # Drop fully-empty trailing cells, keep alignment for the rest.
            cells = list(row)
            while cells and (cells[-1] is None or cells[-1] == ""):
                cells.pop()
            if not cells:
                continue
            rows.append("\t".join("" if c is None else str(c) for c in cells))
        pages.append({"page": i + 1, "text": "\n".join(rows)})
    wb.close()
    return pages or [{"page": 1, "text": ""}]


def _strip_xbrl(html: str) -> str:
    """Strip inline XBRL wrapper tags (ix:...) that bloat SEC filings.
    These wrap every number/text in verbose XML but contain the same content."""
    import re
    # Remove ix: opening tags but keep content: <ix:nonFraction ...>123</ix:nonFraction> → 123
    html = re.sub(r'<ix:[^>]*>', '', html)
    html = re.sub(r'</ix:[^>]*>', '', html)
    # Remove xbrli/link/context blocks (XBRL metadata, not visible content)
    html = re.sub(r'<xbrli:context[^>]*>.*?</xbrli:context>', '', html, flags=re.DOTALL)
    html = re.sub(r'<link:[^>]*>.*?</link:[^>]*>', '', html, flags=re.DOTALL)
    html = re.sub(r'<xbrli:[^>]*>.*?</xbrli:[^>]*>', '', html, flags=re.DOTALL)
    # Remove XML namespace declarations that bloat the file
    html = re.sub(r'\s+xmlns:[a-z\-]+="[^"]*"', '', html)
    return html


def _extract_sec_sections(html: str) -> dict:
    """Split SEC filing HTML into known sections for targeted extraction.
    Returns dict with keys like 'financial_statements', 'mdna', 'risk_factors', etc."""
    import re
    sections = {}
    # SEC filings use Part I/II and Item N headings
    # Find financial statements section (Item 1 in 10-Q, Item 8 in 10-K)
    for pattern in [
        r'(?i)(?:Item\s+1[.\s]*Financial\s+Statements)(.*?)(?=Item\s+[2-9]|Part\s+II|$)',
        r'(?i)(?:Item\s+8[.\s]*Financial\s+Statements)(.*?)(?=Item\s+9|Part\s+[IVX]|$)',
        r'(?i)(?:CONSOLIDATED\s+(?:STATEMENTS?|BALANCE))(.*?)(?=Item\s+\d|Part\s+[IVX]|SIGNATURES|$)',
    ]:
        m = re.search(pattern, html, re.DOTALL)
        if m and len(m.group(1)) > 1000:
            sections['financial_statements'] = m.group(1)
            break
    # MD&A section
    for pattern in [
        r'(?i)(?:Item\s+2[.\s]*Management.s\s+Discussion)(.*?)(?=Item\s+[3-9]|Part\s+II|$)',
        r'(?i)(?:Item\s+7[.\s]*Management.s\s+Discussion)(.*?)(?=Item\s+[89]|Part\s+[IVX]|$)',
    ]:
        m = re.search(pattern, html, re.DOTALL)
        if m and len(m.group(1)) > 500:
            sections['mdna'] = m.group(1)
            break
    return sections


def extract_text_html(file_path: str) -> tuple[list[dict], list[dict]]:
    """
    Extract text and tables from an HTML file (e.g. SEC EDGAR .htm filings).
    Optimised for large SEC filings (10-K can be 8MB+):
    1. Strip XBRL tags first (reduces size 50-80%)
    2. Extract financial sections only for very large files
    3. Parse tables via regex for speed, BeautifulSoup for text
    """
    from bs4 import BeautifulSoup
    import re

    raw = Path(file_path).read_text(errors="replace")
    file_size = len(raw)
    logger.info("HTML extraction starting: %d bytes from %s", file_size, file_path)

    # Step 1: Strip XBRL if present (reduces 8MB → ~2MB typically)
    if 'xmlns:ix=' in raw[:5000] or '<ix:' in raw[:5000]:
        raw = _strip_xbrl(raw)
        logger.info("Stripped XBRL: %d → %d bytes (%.0f%% reduction)",
                     file_size, len(raw), (1 - len(raw)/file_size) * 100)

    # Step 2: For very large files, try targeted section extraction to
    # keep LLM token cost bounded. On modern SEC iXBRL filings the Item
    # N headings are wrapped across tags and styled as divs, which breaks
    # the naive regex — observed on ARW US 10-K 2025_Q4 returning ~2KB
    # of boilerplate instead of the financial statements. When the
    # section extraction produces a clearly-too-small result, fall
    # through to the full stripped content so downstream extraction has
    # something real to work with.
    working_html = raw
    if len(raw) > 2_000_000:  # > 2MB after XBRL strip
        sections = _extract_sec_sections(raw)
        candidate = ""
        if sections:
            candidate = sections.get('financial_statements', '')
            if 'mdna' in sections:
                candidate += '\n\n' + sections['mdna']
            logger.info("Large file: section-extracted %d chars from %d sections",
                         len(candidate), len(sections))

        # A usable financial extract from a 10-K/10-Q should be at LEAST
        # ~50KB (financial statements + MD&A are hundreds of thousands of
        # chars in practice). Anything below that is the iXBRL misfire —
        # fall back to the full document (capped at 2MB) so Claude
        # actually sees the numbers.
        MIN_USEFUL = 50_000
        if len(candidate) >= MIN_USEFUL:
            working_html = candidate
        else:
            working_html = raw[:2_000_000]
            logger.warning(
                "Section extraction produced only %d chars (<%d threshold) — "
                "likely iXBRL misfire. Using full stripped HTML (%d chars).",
                len(candidate), MIN_USEFUL, len(working_html),
            )

    # Step 3: Extract tables via regex (faster than BeautifulSoup for large files)
    tables = []
    table_matches = list(re.finditer(r'<table[^>]*>(.*?)</table>', working_html, re.DOTALL | re.IGNORECASE))
    for i, m in enumerate(table_matches):
        table_html = m.group(0)
        rows = []
        for tr_match in re.finditer(r'<tr[^>]*>(.*?)</tr>', table_html, re.DOTALL | re.IGNORECASE):
            cells = re.findall(r'<t[dh][^>]*>(.*?)</t[dh]>', tr_match.group(1), re.DOTALL | re.IGNORECASE)
            # Strip HTML from cell content
            cleaned_cells = [re.sub(r'<[^>]+>', '', c).strip() for c in cells]
            if any(c for c in cleaned_cells):
                rows.append(cleaned_cells)
        if rows and len(rows) > 1:
            tables.append({"page": i + 1, "tables": [rows]})

    # Step 4: Extract text — try BeautifulSoup, fall back to regex strip
    try:
        soup = BeautifulSoup(working_html, "html.parser")  # html.parser handles partial HTML better than lxml
        for tag in soup.find_all(["script", "style", "meta", "link", "head"]):
            tag.decompose()
        text = soup.get_text(separator="\n")
    except Exception as e:
        logger.warning("BeautifulSoup failed, using regex text extraction: %s", str(e)[:100])
        text = ""

    # Fallback: if BeautifulSoup produced nothing, strip HTML tags via regex
    if len(text.strip()) < 100:
        logger.info("BeautifulSoup produced %d chars, falling back to regex strip", len(text.strip()))
        text = re.sub(r'<[^>]+>', ' ', working_html)
        text = re.sub(r'&nbsp;|&amp;|&lt;|&gt;|&#\d+;|&#x[0-9a-fA-F]+;', ' ', text)
        text = re.sub(r'\s+', ' ', text)

    lines = [line.strip() for line in text.splitlines()]
    cleaned = []
    prev_empty = False
    for line in lines:
        if not line:
            if not prev_empty:
                cleaned.append("")
            prev_empty = True
        else:
            cleaned.append(line)
            prev_empty = False
    full_text = "\n".join(cleaned).strip()

    # Remove SEC boilerplate
    full_text = re.sub(r'(?i)^\s*UNITED STATES\s*\n\s*SECURITIES AND EXCHANGE COMMISSION.*?FORM\s+\d+-[A-Z]+\s*\n', '', full_text, flags=re.DOTALL)

    # Split into page-like chunks
    CHUNK_SIZE = 3000
    pages = []
    pos = 0
    while pos < len(full_text):
        chunk = full_text[pos:pos + CHUNK_SIZE]
        if pos + CHUNK_SIZE < len(full_text):
            last_break = chunk.rfind("\n\n")
            if last_break > CHUNK_SIZE * 0.5:
                chunk = full_text[pos:pos + last_break]
        pages.append({"page": len(pages) + 1, "text": chunk.strip()})
        pos += len(chunk)

    if not pages:
        pages = [{"page": 1, "text": full_text}]

    logger.info("HTML extraction: %d chars, %d pages, %d tables from %s (original %d bytes)",
                len(full_text), len(pages), len(tables), file_path, file_size)
    return pages, tables


# ─────────────────────────────────────────────────────────────────
# Document classification (LLM)
# ─────────────────────────────────────────────────────────────────

def classify_document(text_preview: str, ticker: str = None) -> ClassifiedDocument:
    """Classify a document using the first ~2000 chars."""
    prompt = DOCUMENT_CLASSIFIER.format(text=text_preview[:2000])
    data = call_llm_json(prompt, feature="classification", model="claude-haiku-4-5-20251001", ticker=ticker)
    return ClassifiedDocument(**data)


# ─────────────────────────────────────────────────────────────────
# Native Claude PDF fallback (Tier 1.3)
# ─────────────────────────────────────────────────────────────────

async def _maybe_native_pdf_fallback(
    *, file_path: str, ext: str, tables: list, doc_type: str,
) -> list:
    """Decide + run the native-PDF fallback.

    Triggers when:
      - feature flag `settings.native_pdf_fallback` is on
      - file is a PDF
      - baseline produced 0 tables (or exactly 0 after pdfplumber timeout)
      - doc type is one we expect tables on (SEC 10-Q/10-K, Canadian
        condensed financial statements, etc.)

    Returns either the original `tables` or the fallback-produced
    tables coerced into pdfplumber's shape. Never raises — a fallback
    failure degrades silently to the baseline result.
    """
    if not getattr(settings, "native_pdf_fallback", False):
        return tables
    if ext != ".pdf":
        return tables
    if tables:
        # Baseline found something — trust it.
        return tables

    from services.native_pdf_fallback import (
        extract_tables_from_pdf, tables_to_baseline_shape, DEFAULT_FALLBACK_DOC_TYPES,
    )
    if doc_type not in DEFAULT_FALLBACK_DOC_TYPES:
        logger.info(
            "native_pdf_fallback: skipped — doc_type=%r not in fallback set",
            doc_type,
        )
        return tables

    logger.info("native_pdf_fallback: triggered for %s (type=%s)",
                Path(file_path).name, doc_type)
    try:
        result = await extract_tables_from_pdf(file_path)
    except Exception as exc:  # noqa: BLE001
        logger.warning("native_pdf_fallback raised: %s", str(exc)[:200])
        return tables
    if result.get("status") != "ok" or not result.get("tables"):
        logger.info("native_pdf_fallback: no tables recovered (%s)",
                    result.get("reason", "")[:150])
        return tables
    recovered = result["tables"]
    logger.info("native_pdf_fallback: recovered %d tables, cost=$%s, retried=%s",
                len(recovered), result.get("cost_usd"), result.get("retried"))
    return tables_to_baseline_shape(recovered)


# ─────────────────────────────────────────────────────────────────
# Full processing pipeline
# ─────────────────────────────────────────────────────────────────

async def process_document(db: AsyncSession, document: Document, ticker: str = "UNKNOWN") -> dict:
    """
    End-to-end processing of a single document:
      1. Extract text + tables
      2. Classify
      3. Split into sections and persist
      4. Write processed JSON files
    Returns a summary dict.
    """
    file_path = document.file_path
    ext = Path(file_path).suffix.lower()

    # Raw file is a cache — services.doc_fetch.ensure_local_file handles
    # lazy re-download from source_url when the local copy is missing
    # (evicted under disk pressure, volume reset, etc). Single code path
    # shared with the /file endpoint.
    from services.doc_fetch import ensure_local_file
    await ensure_local_file(document)

    # 1. Extract
    if ext == ".pdf":
        pages = extract_text_pymupdf(file_path)
        # pdfplumber has been observed hanging for 500+ seconds on
        # Canadian condensed financial statements (Sprint C-prep A/B).
        # Cap it well below that so the native-PDF fallback can fire
        # instead of the pipeline stalling.
        import asyncio as _asyncio
        _timeout = getattr(settings, "native_pdf_baseline_timeout_seconds", 60)
        try:
            tables = await _asyncio.wait_for(
                _asyncio.to_thread(extract_tables_pdfplumber, file_path),
                timeout=_timeout,
            )
        except _asyncio.TimeoutError:
            logger.warning(
                "pdfplumber exceeded %ds on %s — skipping baseline tables, "
                "native-PDF fallback will handle if eligible",
                _timeout, Path(file_path).name,
            )
            tables = []
        except Exception as _exc:
            logger.warning("pdfplumber failed on %s: %s", Path(file_path).name, _exc)
            tables = []
    elif ext in (".docx", ".doc"):
        pages = extract_text_docx(file_path)
        tables = []
    elif ext in (".xlsx", ".xls", ".xlsm"):
        pages = extract_text_xlsx(file_path)
        tables = []
    elif ext in (".htm", ".html", ".xhtml"):
        pages, tables = extract_text_html(file_path)
    else:
        # Plain text fallback
        raw_text = Path(file_path).read_text(errors="replace")
        # Chunk into pages for consistency
        CHUNK = 3000
        pages = [{"page": i + 1, "text": raw_text[i:i+CHUNK].strip()}
                 for i in range(0, len(raw_text), CHUNK)] or [{"page": 1, "text": raw_text}]
        tables = []

    full_text = "\n\n".join(p["text"] for p in pages)

    # Store parsed pages on document object so metric_extractor can use
    # real page boundaries for the pre-filter (avoids 3000-char fallback)
    document._parsed_pages = pages

    # 2. Classify — skip LLM if document_type already known (e.g. from EDGAR form type)
    if document.document_type and document.document_type not in ('other', 'unknown'):
        from schemas import ClassifiedDocument
        classification = ClassifiedDocument(
            document_type=document.document_type,
            company_ticker=ticker,
            period_label=document.period_label,
            title=document.title,
            confidence=1.0,
        )
        logger.info("Skipping LLM classification — type already known: %s", document.document_type)
    else:
        classification = classify_document(full_text, ticker=ticker)

    # 2b. Native-PDF fallback (Tier 1.3) — fires when baseline returned
    # zero tables on a doc type that should have them. Silently noops
    # when `settings.native_pdf_fallback` is False (default during
    # shadow-run period).
    tables = await _maybe_native_pdf_fallback(
        file_path=file_path, ext=ext, tables=tables,
        doc_type=(classification.document_type or document.document_type or "").lower(),
    )

    # 3. Persist sections (strip null bytes — PostgreSQL rejects \x00 in text)
    section_objs: list[DocumentSection] = []
    for p in pages:
        clean_text = p["text"].replace("\x00", "") if p.get("text") else ""
        p["text"] = clean_text  # also clean for JSON export
        section = DocumentSection(
            id=uuid.uuid4(),
            document_id=document.id,
            section_type="page",
            section_title=f"Page {p['page']}",
            page_number=p["page"],
            text_content=clean_text,
        )
        db.add(section)
        section_objs.append(section)

    # 3b. Tier 3.4 — embed each section for semantic search. Gated on
    # settings.use_pgvector_search so we don't pay the ~130MB model
    # cold-start when the feature is off. Embedding failure is
    # non-fatal — sections still get written, just without .embedding,
    # and the UI falls back to keyword search for those rows.
    if settings.use_pgvector_search and section_objs:
        try:
            from services.vector_search import embed_texts
            vecs = await embed_texts([s.text_content or "" for s in section_objs])
            for sec, vec in zip(section_objs, vecs):
                if vec is not None:
                    sec.embedding = vec
            _ok = sum(1 for v in vecs if v is not None)
            logger.info("Embedded %d/%d sections for doc %s", _ok, len(vecs), document.id)
        except Exception as exc:
            logger.warning("Section embedding failed for doc %s: %s", document.id, str(exc)[:200])

    # 4. Write processed JSON files
    proc_dir = Path(settings.storage_base_path) / "processed" / ticker / (document.period_label or "misc")
    proc_dir.mkdir(parents=True, exist_ok=True)

    (proc_dir / "parsed_text.json").write_text(json.dumps(pages, indent=2))
    (proc_dir / "tables.json").write_text(json.dumps(tables, indent=2))
    (proc_dir / "metadata.json").write_text(classification.model_dump_json(indent=2))

    # 5. Update status
    document.parsing_status = "completed"
    if classification.document_type and not document.document_type:
        document.document_type = classification.document_type
    await db.commit()

    logger.info("Processed document %s → %s", document.id, classification.document_type)
    return {
        "document_id": str(document.id),
        "pages": len(pages),
        "tables_found": sum(len(t.get("tables", [])) for t in tables),
        "classification": classification.model_dump(),
        "full_text": full_text,  # Return directly for parallel processing
        "tables_data": tables,
    }
